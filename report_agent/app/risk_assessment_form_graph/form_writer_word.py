from __future__ import annotations

import copy
import io
from datetime import date
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.table import Table, _Cell

from app.risk_assessment_form_graph.schemas import FinalHistoryRow, RiskDataCorrectionResult

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_FORM_PATH =  PROJECT_ROOT / "report_template" / "위험성평가표.docx"
DEFAULT_OUTPUT_PATH = PROJECT_ROOT/"risk_assessment_form"/"risk_assessment_form_filled.docx"

META_ROW_INDEX = 2
COMPANY_NAME_TC_INDEX = 1
GENERATED_DATE_TC_INDEX = 3
DATA_START_ROW_INDEX = 7
DATA_COLUMN_COUNT = 16

# 0-based column indices that hold image URLs (조치 전 사진 / 조치 후 사진).
IMAGE_COLUMN_INDEXES = {7, 13}
IMAGE_WIDTH_EMU = 1143000


def _value(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _date_text(value: Any) -> str:
    text = _value(value)
    return text.replace("T", " ") if "T" in text else text


def _first_non_empty(*values: Any) -> str:
    for value in values:
        text = _value(value).strip()
        if text:
            return text
    return ""


def _row_to_dict(row: FinalHistoryRow | dict[str, Any]) -> dict[str, Any]:
    if hasattr(row, "model_dump"):
        return row.model_dump(mode="json")
    return dict(row or {})


def _company_info(source_data: dict[str, Any]) -> dict[str, str]:
    company = source_data.get("company") or {}
    users = source_data.get("user") or []
    safety_manager = ""

    for user in users:
        role = str(user.get("role") or "")
        category = str(user.get("category") or "")
        if "안전" in role or "안전" in category:
            safety_manager = str(user.get("name") or "")
            break

    return {
        "company_name": str(company.get("company_name") or ""),
        "safety_manager": safety_manager,
    }


def _form_row(row: FinalHistoryRow | dict[str, Any]) -> list[str]:
    data = _row_to_dict(row)
    # Order matches the template header: 카테고리/위험도/점검이름/구역/점검일시/담당자/내용/
    # 조치 전 사진/조치이름/구역/조치 일시/담당자/조치 내용/조치 후 사진/승인자/타입
    return [
        _value(data.get("category")),
        _value(data.get("risk")),
        _value(data.get("category_name")),
        _value(data.get("inspection_location")),
        _date_text(data.get("inspection_date")),
        _value(data.get("inspection_user_name")),
        _value(data.get("inspection_content")),
        _value(data.get("inspection_image_url")),
        _value(data.get("action_name")),
        _value(data.get("action_location")),
        _date_text(data.get("completed_at")),
        _value(data.get("handler_name")),
        _value(data.get("content")),
        _value(data.get("image_url")),
        _value(data.get("approver_name")),
        _value(data.get("type")),
    ]


def _download_image(url: str | None) -> bytes | None:
    if not url or not url.startswith(("http://", "https://")):
        return None
    try:
        request = Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urlopen(request, timeout=15) as response:
            return response.read()
    except (HTTPError, URLError, ValueError, TimeoutError):
        return None


def _row_tcs(table: Table, row_index: int):
    return table.rows[row_index]._tr.findall(qn("w:tc"))


def _paragraph_mark_rpr(paragraph):
    # Some template cells hold an empty paragraph whose formatting (font size etc.)
    # lives only on the paragraph-mark rPr, with no actual <w:r> run underneath.
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    return p_pr.find(qn("w:rPr"))


def _new_run_with_paragraph_formatting(paragraph, text: str = ""):
    run = paragraph.add_run(text)
    if run._r.find(qn("w:rPr")) is not None:
        return run
    default_rpr = _paragraph_mark_rpr(paragraph)
    if default_rpr is not None:
        run._r.insert(0, copy.deepcopy(default_rpr))
    return run


def _set_tc_text(tc, table: Table, value: str) -> None:
    # Write through the first existing run so template formatting (font/color) survives;
    # cell.text = value would replace the paragraph and drop that formatting. New runs
    # (for cells with no run at all) borrow the paragraph-mark's rPr for the same reason.
    cell = _Cell(tc, table)
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = value
        return

    first_paragraph = paragraphs[0]
    runs = first_paragraph.runs
    if runs:
        runs[0].text = value
        for extra_run in runs[1:]:
            extra_run.text = ""
    else:
        _new_run_with_paragraph_formatting(first_paragraph, value)

    for extra_paragraph in paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""


def _set_tc_image_or_text(tc, table: Table, value: str) -> None:
    image_bytes = _download_image(value)
    if not image_bytes:
        _set_tc_text(tc, table, value)
        return

    cell = _Cell(tc, table)
    paragraphs = cell.paragraphs
    first_paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()
    for run in list(first_paragraph.runs):
        run.text = ""
    for extra_paragraph in paragraphs[1:]:
        for run in extra_paragraph.runs:
            run.text = ""

    run = (
        first_paragraph.runs[0]
        if first_paragraph.runs
        else _new_run_with_paragraph_formatting(first_paragraph)
    )
    run.add_picture(io.BytesIO(image_bytes), width=Emu(IMAGE_WIDTH_EMU))


def _ensure_row_pair(table: Table, pair_index: int) -> int:
    # Each data row is a vertically merged "restart" + "continue" <w:tr> pair. The
    # template pre-builds a handful of blank pairs; clone the last one for overflow.
    restart_index = DATA_START_ROW_INDEX + pair_index * 2
    while len(table.rows) <= restart_index + 1:
        last_restart = table.rows[-2]._tr
        last_continue = table.rows[-1]._tr
        table._tbl.append(copy.deepcopy(last_restart))
        table._tbl.append(copy.deepcopy(last_continue))
    return restart_index


def _write_data_row(table: Table, restart_index: int, values: list[str]) -> None:
    tcs = _row_tcs(table, restart_index)
    for column_index, value in enumerate(values[:DATA_COLUMN_COUNT]):
        if column_index >= len(tcs):
            continue
        if column_index in IMAGE_COLUMN_INDEXES:
            _set_tc_image_or_text(tcs[column_index], table, value)
        else:
            _set_tc_text(tcs[column_index], table, value)


def fill_risk_assessment_form_docx(
    source_data: dict[str, Any],
    correction_result: RiskDataCorrectionResult,
    form_path: Path | str = DEFAULT_FORM_PATH,
    output_path: Path | str = DEFAULT_OUTPUT_PATH,
) -> str:
    form_path = Path(form_path)
    output_path = Path(output_path)

    document = Document(str(form_path))
    table = document.tables[0]

    info = _company_info(source_data)
    meta_tcs = _row_tcs(table, META_ROW_INDEX)
    company_name = _first_non_empty(info.get("company_name"))
    if company_name and COMPANY_NAME_TC_INDEX < len(meta_tcs):
        _set_tc_text(meta_tcs[COMPANY_NAME_TC_INDEX], table, company_name)
    if GENERATED_DATE_TC_INDEX < len(meta_tcs):
        _set_tc_text(meta_tcs[GENERATED_DATE_TC_INDEX], table, date.today().isoformat())

    for pair_index, row in enumerate(correction_result.corrected_rows):
        restart_index = _ensure_row_pair(table, pair_index)
        _write_data_row(table, restart_index, _form_row(row))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(str(output_path))
        return str(output_path)
    except PermissionError:
        fallback_path = output_path.with_name(f"{output_path.stem}_new{output_path.suffix}")
        document.save(str(fallback_path))
        return str(fallback_path)
