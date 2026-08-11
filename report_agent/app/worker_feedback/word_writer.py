from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

DEFAULT_TEMPLATE_PATH = Path(
    "output/종사자에_의한_유해_위험요인_조사표.docx"
)
DEFAULT_OUTPUT_DIR = Path("output/worker_feedback_reports")

ROW_PLACEHOLDER_MAP = {
    "category": "{{category}}",
    "risk": "{{risk}}",
    "board_created_at": "{{board_created_at}}",
    "board_writer": "{{board_writer}}",
    "category_name": "{{category_name}}",
    "board_contents": "{{board_contents}}",
    "status": "{{status}}",
    "location": "{{location}}",
    "completed_at": "{{completed_at}}",
    "action_status": "{{action_status}}",
    "handler_name": "{{handler_name}}",
    "content": "{{content}}",
    "user": "{{user}}",
}

PLACEHOLDER_ALIASES = {
    "{{board_status}}": "status",
}

IMAGE_PLACEHOLDER_MAP = {
    "board_image_url": "{{board_image_url}}",
    "image_url": "{{image_url}}",
}
IMAGE_MAX_WIDTH_INCHES = 2.2


def _load_document(template_path: Path):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to write docx files. "
            "Install it with: pip install python-docx"
        ) from exc

    return Document(template_path)


def _download_image(url: str | None) -> bytes | None:
    if not url:
        return None
    request = Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urlopen(request, timeout=15) as response:
            return response.read()
    except (HTTPError, URLError, ValueError, TimeoutError):
        return None


def _row_to_dict(row: Any) -> dict[str, Any]:
    if hasattr(row, "model_dump"):
        return row.model_dump(mode="json")
    return dict(row or {})


def _safe_filename_part(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r'[<>:"/\\|?*]', "_", text)
    text = re.sub(r"\s+", "_", text)
    return text[:40] or "worker_feedback"


def _available_output_path(output_path: Path) -> Path:
    if not output_path.exists():
        return output_path

    for index in range(2, 1000):
        candidate = output_path.with_name(
            f"{output_path.stem}_{index:02d}{output_path.suffix}"
        )
        if not candidate.exists():
            return candidate

    raise RuntimeError(f"Could not find available output path for {output_path}")


REPORT_FILENAME_PREFIX = "종사자에 의한 유해 위험요인 보고서_"


def _existing_report_for_board(output_dir: Path, board_id: Any) -> Path | None:
    if board_id is None or not output_dir.exists():
        return None
    matches = sorted(output_dir.glob(f"{REPORT_FILENAME_PREFIX}{board_id}.docx"))
    return matches[0] if matches else None


def _replacement_map(row_data: dict[str, Any]) -> dict[str, str]:
    replacements = {
        placeholder: "" if row_data.get(field) is None else str(row_data.get(field))
        for field, placeholder in ROW_PLACEHOLDER_MAP.items()
    }
    for placeholder, field in PLACEHOLDER_ALIASES.items():
        replacements[placeholder] = (
            "" if row_data.get(field) is None else str(row_data.get(field))
        )
    return replacements


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    changed_in_run = False
    for run in paragraph.runs:
        replaced_text = run.text
        for placeholder, value in replacements.items():
            replaced_text = replaced_text.replace(placeholder, value)
        if replaced_text != run.text:
            run.text = replaced_text
            changed_in_run = True

    if changed_in_run:
        return

    original_text = paragraph.text
    replaced_text = original_text
    for placeholder, value in replacements.items():
        replaced_text = replaced_text.replace(placeholder, value)

    if replaced_text == original_text or not paragraph.runs:
        return

    paragraph.runs[0].text = replaced_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


def _clear_placeholder_text(paragraph, placeholder: str) -> bool:
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, "")
            return True

    if placeholder not in paragraph.text or not paragraph.runs:
        return False

    paragraph.runs[0].text = paragraph.text.replace(placeholder, "")
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _embed_image_in_paragraph(paragraph, placeholder: str, image_bytes: bytes | None) -> None:
    from docx.shared import Inches

    if not _clear_placeholder_text(paragraph, placeholder):
        return
    if not image_bytes:
        return

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    try:
        run.add_picture(BytesIO(image_bytes), width=Inches(IMAGE_MAX_WIDTH_INCHES))
    except Exception:
        pass


def _embed_images_in_table(table, image_map: dict[str, bytes | None]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for placeholder, image_bytes in image_map.items():
                    _embed_image_in_paragraph(paragraph, placeholder, image_bytes)
            for nested_table in cell.tables:
                _embed_images_in_table(nested_table, image_map)


def _embed_row_images(document, row_data: dict[str, Any]) -> None:
    image_map = {
        placeholder: _download_image(row_data.get(field))
        for field, placeholder in IMAGE_PLACEHOLDER_MAP.items()
    }

    for paragraph in document.paragraphs:
        for placeholder, image_bytes in image_map.items():
            _embed_image_in_paragraph(paragraph, placeholder, image_bytes)

    for table in document.tables:
        _embed_images_in_table(table, image_map)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for placeholder, image_bytes in image_map.items():
                _embed_image_in_paragraph(paragraph, placeholder, image_bytes)
        for table in section.header.tables:
            _embed_images_in_table(table, image_map)
        for paragraph in section.footer.paragraphs:
            for placeholder, image_bytes in image_map.items():
                _embed_image_in_paragraph(paragraph, placeholder, image_bytes)
        for table in section.footer.tables:
            _embed_images_in_table(table, image_map)


def _replace_placeholders(document, row_data: dict[str, Any]) -> None:
    replacements = _replacement_map(row_data)

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        _replace_in_table(table, replacements)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            _replace_in_table(table, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            _replace_in_table(table, replacements)


def fill_worker_feedback_word_report(
    row: Any,
    output_path: Path | str,
    template_path: Path | str = DEFAULT_TEMPLATE_PATH,
) -> str:
    template_path = Path(template_path)
    output_path = Path(output_path)
    row_data = _row_to_dict(row)

    document = _load_document(template_path)
    _replace_placeholders(document, row_data)
    _embed_row_images(document, row_data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path = _available_output_path(output_path)
    document.save(output_path)
    return str(output_path)


def fill_worker_feedback_word_reports(
    corrected_rows: list[Any],
    template_path: Path | str = DEFAULT_TEMPLATE_PATH,
    output_dir: Path | str = DEFAULT_OUTPUT_DIR,
) -> list[str]:
    output_dir = Path(output_dir)
    output_paths: list[str] = []

    for index, row in enumerate(corrected_rows, start=1):
        row_data = _row_to_dict(row)
        board_id = row_data.get("board_id")

        existing_report_path = _existing_report_for_board(output_dir, board_id)
        if existing_report_path is not None:
            output_paths.append(str(existing_report_path))
            continue

        if board_id is not None:
            output_path = output_dir / f"{REPORT_FILENAME_PREFIX}{board_id}.docx"
        else:
            name_part = _safe_filename_part(
                row_data.get("category_name")
                or row_data.get("category")
                or row_data.get("location")
            )
            output_path = output_dir / f"{REPORT_FILENAME_PREFIX}{index:03d}_{name_part}.docx"
        output_paths.append(
            fill_worker_feedback_word_report(
                row,
                output_path=output_path,
                template_path=template_path,
            )
        )

    return output_paths
